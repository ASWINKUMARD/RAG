from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
from typing import List, Optional, Dict
import logging
import os
import io
import httpx
import re

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

router = APIRouter()

# Global storage for documents and chunks
document_store: Dict[str, List[Dict]] = {}
uploaded_files: List[str] = []

# -----------------------------
# Pydantic Models
# -----------------------------
class ChatMessage(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    message: str = Field(..., min_length=1)
    conversation_history: Optional[List[ChatMessage]] = Field(default_factory=list)

class ChatResponse(BaseModel):
    response: str
    sources: Optional[List[str]] = Field(default_factory=list)
    status: str = "success"

class UploadResponse(BaseModel):
    message: str
    filename: str
    chunks_created: int = 0
    status: str = "success"

# -----------------------------
# Text Extraction Functions
# -----------------------------
def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF with multiple fallback methods"""
    text = ""
    
    # Method 1: PyPDF2 (fastest, try first)
    try:
        import PyPDF2
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        for page_num in range(len(pdf_reader.pages)):
            try:
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    page_text = re.sub(r'\s+', ' ', page_text)
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"PyPDF2 page {page_num} error: {e}")
        
        text = '\n\n'.join(text_parts)
        if len(text.strip()) > 100:
            logger.info(f"✓ PyPDF2 extracted {len(text)} chars")
            return text
    except ImportError:
        logger.warning("PyPDF2 not installed")
    except Exception as e:
        logger.warning(f"PyPDF2 failed: {e}")
    
    # Method 2: pdfplumber (more robust for complex PDFs)
    try:
        import pdfplumber
        pdf_file = io.BytesIO(content)
        
        text_parts = []
        with pdfplumber.open(pdf_file) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        page_text = re.sub(r'\s+', ' ', page_text)
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"pdfplumber page {page_num} error: {e}")
        
        text = '\n\n'.join(text_parts)
        if len(text.strip()) > 100:
            logger.info(f"✓ pdfplumber extracted {len(text)} chars")
            return text
    except ImportError:
        logger.info("pdfplumber not available")
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")
    
    return text

def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX"""
    try:
        import docx
        
        docx_file = io.BytesIO(content)
        doc = docx.Document(docx_file)
        
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = '\n\n'.join(text_parts)
        
        if len(full_text.strip()) > 20:
            logger.info(f"✓ Extracted {len(full_text)} chars from DOCX")
            return full_text
        return ""
            
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx not installed. Install with: pip install python-docx"
        )
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""

def extract_text_from_file(content: bytes, filename: str) -> str:
    """Main text extraction dispatcher"""
    try:
        file_lower = filename.lower()
        
        if file_lower.endswith('.txt'):
            # Try multiple encodings for text files
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = content.decode(encoding)
                    if len(text.strip()) > 20:
                        logger.info(f"✓ TXT decoded with {encoding}: {len(text)} chars")
                        return text
                except (UnicodeDecodeError, AttributeError):
                    continue
            return ""
        
        elif file_lower.endswith('.pdf'):
            text = extract_text_from_pdf(content)
            return text.strip()
        
        elif file_lower.endswith('.docx'):
            return extract_text_from_docx(content)
        
        else:
            # Try to decode as UTF-8 for unknown types
            return content.decode('utf-8', errors='ignore')
            
    except Exception as e:
        logger.error(f"Extraction error for {filename}: {e}")
        return ""

# -----------------------------
# Chunking Functions
# -----------------------------
def create_chunks(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict]:
    """Create overlapping chunks with metadata"""
    
    # Clean and normalize text
    text = re.sub(r'\s+', ' ', text).strip()
    
    if not text or len(text) < 50:
        logger.warning("Text too short for chunking")
        return []
    
    # Split into sentences
    sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
    
    chunks = []
    current_chunk = []
    current_length = 0
    chunk_idx = 0
    
    for sentence in sentences:
        words = sentence.split()
        word_count = len(words)
        
        # Check if we should create a new chunk
        if current_length + word_count > chunk_size and current_chunk:
            # Save current chunk
            chunk_text = ' '.join(current_chunk)
            
            # Create searchable word set (lowercase, cleaned)
            words_clean = set()
            for word in chunk_text.lower().split():
                cleaned = re.sub(r'[^\w\s]', '', word)
                if cleaned and len(cleaned) > 1:  # Skip single chars
                    words_clean.add(cleaned)
            
            chunks.append({
                'text': chunk_text,
                'index': chunk_idx,
                'words': words_clean,
                'char_count': len(chunk_text),
                'word_count': len(current_chunk)
            })
            
            logger.debug(f"Chunk {chunk_idx}: {len(chunk_text)} chars, {len(current_chunk)} words")
            chunk_idx += 1
            
            # Start new chunk with overlap
            if overlap > 0 and len(current_chunk) > overlap:
                current_chunk = current_chunk[-overlap:] + words
                current_length = len(current_chunk)
            else:
                current_chunk = words
                current_length = word_count
        else:
            current_chunk.extend(words)
            current_length += word_count
    
    # Don't forget the last chunk
    if current_chunk:
        chunk_text = ' '.join(current_chunk)
        words_clean = set()
        for word in chunk_text.lower().split():
            cleaned = re.sub(r'[^\w\s]', '', word)
            if cleaned and len(cleaned) > 1:
                words_clean.add(cleaned)
        
        chunks.append({
            'text': chunk_text,
            'index': chunk_idx,
            'words': words_clean,
            'char_count': len(chunk_text),
            'word_count': len(current_chunk)
        })
        
        logger.debug(f"Chunk {chunk_idx} (final): {len(chunk_text)} chars")
    
    logger.info(f"✓ Created {len(chunks)} chunks from {len(text)} chars")
    return chunks

# -----------------------------
# Retrieval Functions
# -----------------------------
def retrieve_relevant_chunks(query: str, top_k: int = 4) -> List[tuple]:
    """Enhanced retrieval with semantic-aware scoring"""
    
    query_lower = query.lower()
    
    # Clean and extract query words
    query_words = []
    for word in query_lower.split():
        cleaned = re.sub(r'[^\w\s]', '', word)
        if cleaned and len(cleaned) > 1:
            query_words.append(cleaned)
    
    query_words_set = set(query_words)
    
    # Enhanced stop words list
    stop_words = {
        'the', 'is', 'at', 'which', 'on', 'a', 'an', 'and', 'or', 
        'but', 'in', 'with', 'to', 'for', 'of', 'it', 'as', 'by',
        'this', 'that', 'what', 'are', 'be', 'was', 'were', 'been',
        'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would',
        'should', 'could', 'can', 'may', 'might', 'must', 'from'
    }
    
    query_keywords = query_words_set - stop_words
    
    if not query_keywords:
        query_keywords = query_words_set
    
    logger.info(f"Query: '{query}' | Keywords: {query_keywords}")
    
    results = []
    
    for doc_id, chunks in document_store.items():
        logger.debug(f"Searching {doc_id}: {len(chunks)} chunks")
        
        for chunk in chunks:
            chunk_text = chunk['text']
            chunk_lower = chunk_text.lower()
            chunk_words = chunk['words']
            
            score = 0
            
            # 1. EXACT PHRASE MATCH (highest priority)
            if query_lower in chunk_lower:
                score += 1000
                logger.debug(f"✓ Exact phrase match in chunk {chunk['index']}")
            
            # 2. ALL KEYWORDS PRESENT (very high score)
            if query_keywords.issubset(chunk_words):
                score += 500
                logger.debug(f"✓ All keywords present in chunk {chunk['index']}")
            
            # 3. KEYWORD OVERLAP RATIO
            overlap = len(query_keywords & chunk_words)
            if overlap > 0:
                overlap_ratio = overlap / len(query_keywords)
                score += overlap_ratio * 300
                logger.debug(f"Chunk {chunk['index']}: {overlap}/{len(query_keywords)} keywords ({overlap_ratio:.1%})")
            
            # 4. KEYWORD FREQUENCY (TF component)
            for keyword in query_keywords:
                count = chunk_lower.count(keyword)
                if count > 0:
                    keyword_score = min(count * 50, 150)
                    score += keyword_score
            
            # 5. QUERY WORD ORDER BONUS (bigrams and trigrams)
            if len(query_words) > 1:
                # Bigrams
                for i in range(len(query_words) - 1):
                    bigram = f"{query_words[i]} {query_words[i+1]}"
                    if bigram in chunk_lower:
                        score += 200
                
                # Trigrams
                if len(query_words) > 2:
                    for i in range(len(query_words) - 2):
                        trigram = f"{query_words[i]} {query_words[i+1]} {query_words[i+2]}"
                        if trigram in chunk_lower:
                            score += 300
            
            # 6. PROXIMITY BONUS (keywords close together)
            if overlap > 1:
                keyword_positions = []
                for keyword in query_keywords:
                    if keyword in chunk_lower:
                        pos = chunk_lower.find(keyword)
                        keyword_positions.append(pos)
                
                if len(keyword_positions) > 1:
                    keyword_positions.sort()
                    max_distance = keyword_positions[-1] - keyword_positions[0]
                    if max_distance < 100:
                        proximity_score = (100 - max_distance) * 2
                        score += proximity_score
            
            if score > 0:
                results.append((score, doc_id, chunk_text, chunk['index']))
    
    # Sort by score (highest first)
    results.sort(reverse=True, key=lambda x: x[0])
    
    if results:
        logger.info(f"✓ Found {len(results)} matching chunks")
        logger.info(f"Top scores: {[f'{r[0]:.0f}' for r in results[:5]]}")
    else:
        logger.warning("⚠ No matching chunks found!")
    
    return results[:top_k]

# -----------------------------
# LLM Generation
# -----------------------------
async def generate_llm_response(query: str, context: str, sources: List[str]) -> str:
    """Generate AI response using OpenRouter"""
    
    system_prompt = """You are a precise RAG (Retrieval-Augmented Generation) assistant. Your task is to answer questions using ONLY the information provided in the context below.

CRITICAL INSTRUCTIONS:
1. Answer ONLY using information from the provided context
2. Be specific, detailed, and well-organized in your responses
3. Use bullet points or numbered lists when presenting multiple items
4. Quote relevant parts of the context when appropriate
5. If the context contains the answer, provide it completely and accurately
6. If the context does NOT contain enough information, clearly state what you know and what's missing
7. NEVER invent, assume, or add information not present in the context
8. Maintain a helpful and professional tone"""

    user_prompt = f"""CONTEXT FROM UPLOADED DOCUMENTS:
{context}

---

USER QUESTION: {query}

Please answer the user's question based ONLY on the context provided above. Be thorough and cite specific details from the context."""

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    
    try:
        api_key = os.getenv('OPENROUTER_API_KEY', '')
        
        if not api_key:
            logger.warning("No OPENROUTER_API_KEY set - using fallback response")
            return format_fallback_response(query, context, sources)
        
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                    "HTTP-Referer": "https://rag-chatbot.com",
                    "X-Title": "RAG Chatbot"
                },
                json={
                    "model": "kwaipilot/kat-coder-pro:free",
                    "messages": messages,
                    "temperature": 0.1,
                    "max_tokens": 2000,
                    "top_p": 0.95
                }
            )
            
            if response.status_code == 200:
                result = response.json()
                answer = result['choices'][0]['message']['content']
                logger.info(f"✓ LLM generated {len(answer)} char response")
                return answer
            else:
                logger.error(f"OpenRouter API error: {response.status_code} - {response.text}")
                return format_fallback_response(query, context, sources)
                
    except httpx.TimeoutException:
        logger.error("OpenRouter API timeout")
        return format_fallback_response(query, context, sources)
    except Exception as e:
        logger.error(f"LLM generation error: {e}")
        return format_fallback_response(query, context, sources)

def format_fallback_response(query: str, context: str, sources: List[str]) -> str:
    """Format a structured response when LLM is not available"""
    
    response = f"📄 **Based on the uploaded documents** ({', '.join(sources)}):\n\n"
    
    # Extract and format relevant snippets
    lines = [line.strip() for line in context.split('\n') if line.strip()]
    
    # Show first 20 meaningful lines
    shown = 0
    for line in lines:
        if shown >= 20:
            break
        if len(line) > 30:  # Skip very short lines
            response += f"• {line}\n"
            shown += 1
    
    if len(context) > 1500:
        response += f"\n...(showing excerpt from {len(context)} total characters)\n"
    
    response += "\n\n💡 **Tip**: Set OPENROUTER_API_KEY environment variable for AI-generated answers."
    
    return response

# -----------------------------
# API ENDPOINTS
# -----------------------------
@router.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    """Enhanced chat endpoint with RAG"""
    try:
        query = request.message.strip()
        logger.info(f"\n{'='*60}\nNEW QUERY: {query}\n{'='*60}")
        
        if not document_store:
            return ChatResponse(
                response="📤 **Please upload documents first!**\n\nUse the 'Upload Document' button above to upload PDF, DOCX, or TXT files. Once uploaded, I can answer questions about their content.",
                sources=[],
                status="success"
            )
        
        # Retrieve relevant chunks
        retrieved = retrieve_relevant_chunks(query, top_k=4)
        
        if not retrieved:
            available_docs = ', '.join(list(document_store.keys()))
            return ChatResponse(
                response=f"❌ **No relevant information found**\n\n"
                        f"I couldn't find information about '{query}' in the uploaded documents.\n\n"
                        f"📁 **Available documents**: {available_docs}\n\n"
                        f"💡 **Suggestions**:\n"
                        f"• Try rephrasing your question\n"
                        f"• Use different keywords\n"
                        f"• Ask about topics explicitly covered in your documents\n"
                        f"• Upload additional relevant documents",
                sources=list(document_store.keys()),
                status="success"
            )
        
        # Build context from retrieved chunks
        context_parts = []
        sources = set()
        
        for score, doc_id, chunk_text, chunk_idx in retrieved:
            sources.add(doc_id)
            
            context_section = f"""[Document: {doc_id}]
[Section: {chunk_idx + 1}]
[Relevance Score: {score:.0f}]

{chunk_text}""".strip()
            
            context_parts.append(context_section)
            logger.info(f"Using chunk {chunk_idx} from {doc_id} (score: {score:.1f})")
        
        context = "\n\n--- NEXT SECTION ---\n\n".join(context_parts)
        logger.info(f"Total context: {len(context)} chars from {len(retrieved)} chunks")
        
        # Generate response
        response_text = await generate_llm_response(query, context, list(sources))
        
        return ChatResponse(
            response=response_text,
            sources=list(sources),
            status="success"
        )
        
    except Exception as e:
        logger.error(f"Chat error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/upload", response_model=UploadResponse)
async def upload_document(file: UploadFile = File(...)):
    """Upload and process document"""
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        # Validate file type
        allowed = ['.pdf', '.docx', '.txt']
        file_ext = '.' + file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
        
        if file_ext not in allowed:
            raise HTTPException(
                status_code=400,
                detail=f"File type '{file_ext}' not supported. Allowed: {', '.join(allowed)}"
            )
        
        # Read file content
        content = await file.read()
        
        # Validate file size
        if len(content) > 100 * 1024 * 1024:  # 100MB
            raise HTTPException(status_code=400, detail="File exceeds 100MB limit")
        
        if len(content) < 100:
            raise HTTPException(status_code=400, detail="File too small or empty")
        
        logger.info(f"\n{'='*60}\nUPLOADING: {file.filename} ({len(content)} bytes)\n{'='*60}")
        
        # Extract text
        text = extract_text_from_file(content, file.filename)
        
        if not text or len(text.strip()) < 50:
            raise HTTPException(
                status_code=400,
                detail=f"Could not extract text from '{file.filename}'. "
                       f"Possible reasons: scanned PDF, corrupted file, or password-protected document."
            )
        
        logger.info(f"✓ Extracted {len(text)} characters")
        
        # Create chunks
        chunks = create_chunks(text, chunk_size=1000, overlap=200)
        
        if not chunks:
            raise HTTPException(
                status_code=400,
                detail="Text extracted but chunking failed (text may be too short)"
            )
        
        # Store chunks
        doc_id = file.filename
        document_store[doc_id] = chunks
        
        if doc_id not in uploaded_files:
            uploaded_files.append(doc_id)
        
        logger.info(f"✓ SUCCESS: Stored {len(chunks)} chunks for '{doc_id}'")
        
        return UploadResponse(
            message=f"✅ Successfully processed '{file.filename}'\n\nCreated {len(chunks)} searchable chunks ({len(text):,} total characters)",
            filename=file.filename,
            chunks_created=len(chunks),
            status="success"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@router.get("/status")
async def get_status():
    """Get system status"""
    total_chunks = sum(len(chunks) for chunks in document_store.values())
    
    doc_details = []
    for doc_id, chunks in document_store.items():
        total_chars = sum(c['char_count'] for c in chunks)
        total_words = sum(c['word_count'] for c in chunks)
        doc_details.append({
            "filename": doc_id,
            "chunks": len(chunks),
            "total_chars": total_chars,
            "total_words": total_words
        })
    
    return JSONResponse({
        "status": "online",
        "documents_loaded": len(document_store),
        "total_chunks": total_chunks,
        "documents": doc_details,
        "model": "kwaipilot/kat-coder-pro:free",
        "openrouter_configured": bool(os.getenv('OPENROUTER_API_KEY')),
        "chunk_size": 1000,
        "chunk_overlap": 200
    })

@router.get("/documents")
async def list_documents():
    """List all uploaded documents with statistics"""
    docs = []
    for doc_id in uploaded_files:
        if doc_id in document_store:
            chunks = document_store[doc_id]
            total_chars = sum(c['char_count'] for c in chunks)
            total_words = sum(c['word_count'] for c in chunks)
            docs.append({
                "filename": doc_id,
                "chunks": len(chunks),
                "total_characters": total_chars,
                "total_words": total_words,
                "avg_chunk_size": total_chars // len(chunks) if chunks else 0
            })
    
    return JSONResponse({"documents": docs, "count": len(docs)})

@router.delete("/clear")
async def clear_all():
    """Clear all uploaded documents"""
    document_store.clear()
    uploaded_files.clear()
    logger.info("✓ All documents cleared")
    return JSONResponse({"message": "All documents cleared successfully", "status": "success"})
