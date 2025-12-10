from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
from typing import List, Optional, Dict
import logging
import os
import io
import httpx
import re

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

router = APIRouter()

document_store: Dict[str, List[Dict]] = {}
uploaded_files = []

class ChatMessage(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    message: str = Field(..., min_length=1)
    conversation_history: Optional[List[ChatMessage]] = Field(default=[])

class ChatResponse(BaseModel):
    response: str
    sources: Optional[List[str]] = Field(default=[])
    status: str = "success"

class UploadResponse(BaseModel):
    message: str
    filename: str
    chunks_created: int = 0
    status: str = "success"

def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF with multiple fallback methods"""
    text = ""
    
    # Method 1: pdfplumber (best for tables and formatting)
    try:
        import pdfplumber
        pdf_file = io.BytesIO(content)
        
        text_parts = []
        with pdfplumber.open(pdf_file) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        # Clean text
                        page_text = re.sub(r'\s+', ' ', page_text)
                        text_parts.append(page_text)
                        logger.info(f"Page {page_num + 1}: {len(page_text)} chars")
                except Exception as e:
                    logger.warning(f"pdfplumber page {page_num} error: {e}")
        
        text = '\n\n'.join(text_parts)
        if len(text.strip()) > 100:
            logger.info(f"✓ pdfplumber extracted {len(text)} chars from {len(text_parts)} pages")
            return text
    except ImportError:
        logger.info("pdfplumber not available")
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")
    
    # Method 2: PyPDF2
    try:
        import PyPDF2
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        for page_num in range(len(pdf_reader.pages)):
            try:
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    page_text = re.sub(r'\s+', ' ', page_text)
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"PyPDF2 page {page_num} error: {e}")
        
        text = '\n\n'.join(text_parts)
        if len(text.strip()) > 100:
            logger.info(f"✓ PyPDF2 extracted {len(text)} chars")
            return text
    except Exception as e:
        logger.warning(f"PyPDF2 failed: {e}")
    
    return text

def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX"""
    try:
        import docx
        
        docx_file = io.BytesIO(content)
        doc = docx.Document(docx_file)
        
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = '\n\n'.join(text_parts)
        
        if len(full_text.strip()) > 20:
            logger.info(f"✓ Extracted {len(full_text)} chars from DOCX")
            return full_text
        return ""
            
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx not installed"
        )
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""

def extract_text_from_file(content: bytes, filename: str) -> str:
    """Main text extraction"""
    try:
        file_lower = filename.lower()
        
        if file_lower.endswith('.txt'):
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = content.decode(encoding)
                    if len(text.strip()) > 20:
                        logger.info(f"✓ TXT decoded: {len(text)} chars")
                        return text
                except:
                    continue
            return ""
        
        elif file_lower.endswith('.pdf'):
            text = extract_text_from_pdf(content)
            return text.strip()
        
        elif file_lower.endswith('.docx'):
            return extract_text_from_docx(content)
        
        else:
            return content.decode('utf-8', errors='ignore')
            
    except Exception as e:
        logger.error(f"Extraction error: {e}")
        return ""

def create_chunks(text: str, chunk_size: int = 800, overlap: int = 200) -> List[Dict]:
    """Create larger, better overlapping chunks"""
    
    # Clean and normalize text
    text = re.sub(r'\s+', ' ', text).strip()
    
    if not text or len(text) < 50:
        logger.warning("Text too short for chunking")
        return []
    
    # Split into sentences
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    chunks = []
    current_chunk = []
    current_length = 0
    chunk_idx = 0
    
    for sentence in sentences:
        words = sentence.split()
        word_count = len(words)
        
        # Check if we should create a new chunk
        if current_length + word_count > chunk_size and current_chunk:
            # Save current chunk
            chunk_text = ' '.join(current_chunk)
            
            # Create searchable word set (lowercase, no punctuation)
            words_clean = set()
            for word in chunk_text.lower().split():
                cleaned = re.sub(r'[^\w\s]', '', word)
                if cleaned:
                    words_clean.add(cleaned)
            
            chunks.append({
                'text': chunk_text,
                'index': chunk_idx,
                'words': words_clean,
                'char_count': len(chunk_text)
            })
            
            logger.info(f"Chunk {chunk_idx}: {len(chunk_text)} chars, {len(current_chunk)} words")
            chunk_idx += 1
            
            # Start new chunk with overlap
            if overlap > 0 and len(current_chunk) > overlap:
                current_chunk = current_chunk[-overlap:] + words
                current_length = len(current_chunk)
            else:
                current_chunk = words
                current_length = word_count
        else:
            current_chunk.extend(words)
            current_length += word_count
    
    # Don't forget the last chunk
    if current_chunk:
        chunk_text = ' '.join(current_chunk)
        words_clean = set()
        for word in chunk_text.lower().split():
            cleaned = re.sub(r'[^\w\s]', '', word)
            if cleaned:
                words_clean.add(cleaned)
        
        chunks.append({
            'text': chunk_text,
            'index': chunk_idx,
            'words': words_clean,
            'char_count': len(chunk_text)
        })
        
        logger.info(f"Chunk {chunk_idx} (final): {len(chunk_text)} chars")
    
    logger.info(f"✓ Created {len(chunks)} chunks from {len(text)} chars")
    return chunks

def retrieve_relevant_chunks(query: str, top_k: int = 3) -> List[tuple]:
    """Improved retrieval with better scoring"""
    
    query_lower = query.lower()
    
    # Clean query words
    query_words = []
    for word in query_lower.split():
        cleaned = re.sub(r'[^\w\s]', '', word)
        if cleaned:
            query_words.append(cleaned)
    
    query_words_set = set(query_words)
    
    # Common stop words
    stop_words = {'the', 'is', 'at', 'which', 'on', 'a', 'an', 'and', 'or', 
                  'but', 'in', 'with', 'to', 'for', 'of', 'it', 'as', 'by',
                  'this', 'that', 'what', 'are', 'be', 'was', 'were'}
    
    query_keywords = query_words_set - stop_words
    
    if not query_keywords:
        query_keywords = query_words_set
    
    logger.info(f"Query keywords: {query_keywords}")
    
    results = []
    
    for doc_id, chunks in document_store.items():
        logger.info(f"Searching {doc_id}: {len(chunks)} chunks")
        
        for chunk in chunks:
            chunk_text = chunk['text']
            chunk_lower = chunk_text.lower()
            chunk_words = chunk['words']
            
            score = 0
            
            # 1. Exact phrase match (very high score)
            if query_lower in chunk_lower:
                score += 500
                logger.info(f"Exact match in chunk {chunk['index']}")
            
            # 2. All keywords present (high score)
            if query_keywords.issubset(chunk_words):
                score += 200
                logger.info(f"All keywords in chunk {chunk['index']}")
            
            # 3. Partial keyword match
            overlap = len(query_keywords & chunk_words)
            if overlap > 0:
                overlap_ratio = overlap / len(query_keywords)
                score += overlap_ratio * 150
                logger.info(f"Chunk {chunk['index']}: {overlap}/{len(query_keywords)} keywords matched")
            
            # 4. Individual keyword frequency
            for keyword in query_keywords:
                count = chunk_lower.count(keyword)
                if count > 0:
                    score += count * 40
                    logger.info(f"Keyword '{keyword}' appears {count}x in chunk {chunk['index']}")
            
            # 5. Query word order bonus (bigrams)
            if len(query_words) > 1:
                for i in range(len(query_words) - 1):
                    bigram = f"{query_words[i]} {query_words[i+1]}"
                    if bigram in chunk_lower:
                        score += 100
            
            if score > 0:
                results.append((score, doc_id, chunk_text, chunk['index']))
    
    # Sort by score (highest first)
    results.sort(reverse=True, key=lambda x: x[0])
    
    if results:
        logger.info(f"✓ Found {len(results)} matching chunks")
        logger.info(f"Top 3 scores: {[r[0] for r in results[:3]]}")
    else:
        logger.warning("No matching chunks found!")
    
    return results[:top_k]

async def generate_llm_response(query: str, context: str) -> str:
    """Generate AI response using OpenRouter"""
    
    system_prompt = """You are a helpful RAG assistant that answers questions based on provided documents.

IMPORTANT RULES:
1. Answer the user's question using ONLY the information in the context below
2. Be specific, accurate, and well-organized
3. Use bullet points or numbered lists for clarity
4. If the context has the answer, provide complete details
5. If the context doesn't fully answer the question, say what you DO know and what's missing
6. Do NOT make up information not in the context"""

    user_prompt = f"""CONTEXT FROM DOCUMENTS:
{context}

USER QUESTION: {query}

Please answer the question based on the context above."""

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    
    try:
        api_key = os.getenv('OPENROUTER_API_KEY', '')
        
        if not api_key:
            logger.warning("No OPENROUTER_API_KEY set")
            # Return a formatted version of the context
            return format_fallback_response(query, context)
        
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                    "HTTP-Referer": "https://rag-chatbot.com",
                    "X-Title": "RAG Chatbot"
                },
                json={
                    "model": "kwaipilot/kat-coder-pro:free",
                    "messages": messages,
                    "temperature": 0.2,
                    "max_tokens": 1500,
                    "top_p": 0.9
                }
            )
            
            if response.status_code == 200:
                result = response.json()
                answer = result['choices'][0]['message']['content']
                logger.info(f"✓ LLM generated {len(answer)} char response")
                return answer
            else:
                logger.error(f"OpenRouter API error: {response.status_code}")
                return format_fallback_response(query, context)
                
    except Exception as e:
        logger.error(f"LLM generation error: {e}")
        return format_fallback_response(query, context)

def format_fallback_response(query: str, context: str) -> str:
    """Format a response when LLM is not available"""
    
    # Extract key information from context
    lines = context.split('\n')
    
    response = f"Based on the uploaded documents, here's what I found:\n\n"
    
    # Add context sections
    for line in lines[:15]:  # First 15 lines
        if line.strip():
            response += f"• {line.strip()}\n"
    
    if len(lines) > 15:
        response += f"\n... (showing excerpt, full context is {len(context)} characters)\n"
    
    response += "\n\n💡 Tip: Set OPENROUTER_API_KEY for AI-generated answers."
    
    return response

@router.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    """Chat endpoint with improved RAG"""
    try:
        query = request.message.strip()
        logger.info(f"=== NEW QUERY: {query} ===")
        
        if not document_store:
            return ChatResponse(
                response="📤 Please upload documents first! Use the 'Upload Document' button above to get started.",
                sources=[],
                status="success"
            )
        
        # Retrieve relevant chunks
        retrieved = retrieve_relevant_chunks(query, top_k=3)
        
        if not retrieved:
            available_docs = ', '.join(list(document_store.keys()))
            return ChatResponse(
                response=f"❌ I couldn't find relevant information about '{query}' in the uploaded documents.\n\n"
                        f"📁 Available documents: {available_docs}\n\n"
                        f"💡 Try: rephrasing your question, using different keywords, or asking about topics covered in the documents.",
                sources=list(document_store.keys()),
                status="success"
            )
        
        # Build context from retrieved chunks
        context_parts = []
        sources = set()
        
        for score, doc_id, chunk_text, chunk_idx in retrieved:
            context_parts.append(f"[From: {doc_id}, Section {chunk_idx + 1}, Relevance: {score:.0f}]\n{chunk_text}")
            sources.add(doc_id)
            logger.info(f"Using chunk {chunk_idx} from {doc_id} (score: {score:.1f})")
        
        context = "\n\n" + "="*50 + "\n\n".join(context_parts)
        
        logger.info(f"Context size: {len(context)} chars from {len(retrieved)} chunks")
        
        # Generate response
        response_text = await generate_llm_response(query, context)
        
        return ChatResponse(
            response=response_text,
            sources=list(sources),
            status="success"
        )
        
    except Exception as e:
        logger.error(f"Chat error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/upload", response_model=UploadResponse)
async def upload_document(file: UploadFile = File(...)):
    """Upload document with detailed logging"""
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        allowed = ['.pdf', '.docx', '.txt']
        ext = '.' + file.filename.rsplit('.', 1)[-1].lower()
        
        if ext not in allowed:
            raise HTTPException(
                status_code=400,
                detail=f"File type '{ext}' not supported. Allowed: PDF, DOCX, TXT"
            )
        
        content = await file.read()
        
        if len(content) > 100 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File exceeds 100MB")
        
        if len(content) < 100:
            raise HTTPException(status_code=400, detail="File too small or empty")
        
        logger.info(f"=== UPLOADING: {file.filename} ({len(content)} bytes) ===")
        
        # Extract text
        text = extract_text_from_file(content, file.filename)
        
        if not text or len(text.strip()) < 50:
            raise HTTPException(
                status_code=400,
                detail=f"❌ Could not extract sufficient text from '{file.filename}'. "
                       f"Reasons: (1) PDF contains only scanned images, "
                       f"(2) File is corrupted, (3) File is password-protected. "
                       f"Try converting scanned PDFs with OCR first."
            )
        
        logger.info(f"✓ Extracted {len(text)} characters")
        logger.info(f"First 200 chars: {text[:200]}")
        
        # Create chunks with larger size for better context
        chunks = create_chunks(text, chunk_size=800, overlap=200)
        
        if not chunks:
            raise HTTPException(
                status_code=400, 
                detail="Text extracted but chunking failed"
            )
        
        # Store chunks
        doc_id = file.filename
        document_store[doc_id] = chunks
        
        if doc_id not in uploaded_files:
            uploaded_files.append(doc_id)
        
        logger.info(f"✓ SUCCESS: Stored {len(chunks)} chunks for '{doc_id}'")
        
        return UploadResponse(
            message=f"✅ Successfully processed '{file.filename}' into {len(chunks)} searchable chunks",
            filename=file.filename,
            chunks_created=len(chunks),
            status="success"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@router.get("/status")
async def get_status():
    """Detailed status endpoint"""
    total_chunks = sum(len(chunks) for chunks in document_store.values())
    
    doc_details = []
    for doc_id, chunks in document_store.items():
        total_chars = sum(c['char_count'] for c in chunks)
        doc_details.append({
            "filename": doc_id,
            "chunks": len(chunks),
            "total_chars": total_chars
        })
    
    return JSONResponse({
        "status": "online",
        "documents_loaded": len(document_store),
        "total_chunks": total_chunks,
        "documents": doc_details,
        "model": "kwaipilot/kat-coder-pro:free",
        "openrouter_configured": bool(os.getenv('OPENROUTER_API_KEY'))
    })

@router.get("/documents")
async def list_documents():
    """List documents with details"""
    docs = []
    for doc_id in uploaded_files:
        if doc_id in document_store:
            chunks = document_store[doc_id]
            total_chars = sum(c['char_count'] for c in chunks)
            docs.append({
                "filename": doc_id,
                "chunks": len(chunks),
                "total_characters": total_chars
            })
    
    return JSONResponse({"documents": docs, "count": len(docs)})

@router.delete("/clear")
async def clear_all():
    """Clear all documents"""
    document_store.clear()
    uploaded_files.clear()
    logger.info("✓ All documents cleared")
    return JSONResponse({"message": "All documents cleared", "status": "success"})
